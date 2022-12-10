import docx
import datetime

class Docxcreating:
    def __init__(self, based_id, database):
        self.based_id = based_id
        self.database = database
        self.docxname = self.create_name()
        self.create_docx()

    def create_name(self):
        name = str(datetime.datetime.now())
        name = name.replace(' ', '').replace('-','').replace(':', '')[:14]
        return name

    def parse_base(self):
        result = self.database.get_persons(0, 10)
        for row_number, row_data in enumerate(result):
            if row_number == self.based_id: 
                return row_data

    def create_docx(self):
        doc = docx.Document()
        data = self.parse_base()
        doc.add_paragraph(f"Фамилия: {data[1]}")
        doc.add_paragraph(f"Имя: {data[2]}")
        doc.add_paragraph(f"Отчество: {data[3]}")
        doc.add_paragraph(f"Дата рождения: {data[4]}")
        doc.add_paragraph(f"Место рождения: {data[5]}")
        doc.add_paragraph(f"Место регистрации: {data[6]}")
        doc.add_paragraph(f"Серия и Номер: {data[7]}")
        doc.add_paragraph(f"Кем выдан: {data[8]}")
        doc.add_paragraph(f"Дата выдачи: {data[9]}")
        doc.add_paragraph(f"ИНН: {data[10]}")
        doc.add_paragraph(f"Снилс: {data[11]}")
        doc.save(f"{self.docxname}.docx")




